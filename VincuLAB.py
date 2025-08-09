import pandas as pd
import os
import tkinter as tk
from tkinter import messagebox, ttk, filedialog, simpledialog
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from wordcloud import WordCloud
from thefuzz import process
from datetime import date, datetime
import re
import webbrowser

plt.switch_backend('Agg')  # Para evitar problemas con la interfaz gráfica

class UnifiedReportApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema Unificado de Gestión de Reportes - Educación Especial")
        self.root.geometry("600x500")
        
        # Set default directory to Desktop
        self.directorio_destino = os.path.join(os.path.expanduser('~'), 'Desktop')
        
        # URLs de datos
        self.url_datos = "https://docs.google.com/spreadsheets/d/1p42nIbj66UIn-kyZQ1Ilbx13nxiWKIfEMbcrYMFae84/export?format=xlsx"
        self.url_beneficiarios = "https://docs.google.com/spreadsheets/d/15BR53PUapEaKiz2LYHK8l46R7HNYrRHhdwXREIv9Woo/export?format=csv"
        self.url_ubicacion = "https://docs.google.com/spreadsheets/d/1Vbkt7BkHB4wXJu5iZcHnCjNFcbjWjOCi/export?format=xlsx"
        self.url_encuesta = 'https://docs.google.com/spreadsheets/d/1yi99P2uGsbWuk4X0p1EfF7QG8mgU4NICiR2r_EQ57QE/export?format=csv'
        self.url_password = "https://docs.google.com/spreadsheets/d/1Cya9M8otF1sqTh7lp6BEoRyOpekZp3r8oDjbpw56de0/export?format=csv"
        
        self.create_widgets()
        
    def create_widgets(self):
        # Title
        title_label = tk.Label(self.root, text="Sistema Unificado de Gestión de Reportes", 
                             font=("Arial", 16, "bold"))
        title_label.pack(pady=10)
        
        # Period Selection
        period_frame = tk.Frame(self.root)
        period_frame.pack(pady=10)
        tk.Label(period_frame, text="Período (ejemplo: 2025-1):").pack(side=tk.LEFT)
        self.period_entry = tk.Entry(period_frame, width=15)
        self.period_entry.pack(side=tk.LEFT, padx=5)
        
        # Folder Selection
        folder_frame = tk.Frame(self.root)
        folder_frame.pack(pady=10)
        tk.Label(folder_frame, text="Carpeta de destino:").pack()
        self.folder_button = tk.Button(folder_frame, text="Cambiar Carpeta", command=self.select_folder)
        self.folder_button.pack()
        self.folder_label = tk.Label(folder_frame, text=f"Escritorio: {self.directorio_destino}", 
                                   fg="blue", wraplength=500)
        self.folder_label.pack()
        
        # Main Functions Frame
        main_frame = tk.Frame(self.root)
        main_frame.pack(pady=20, expand=True, fill='both')
        
        # Report Generation Section
        report_frame = tk.LabelFrame(main_frame, text="Generación de Reportes", padx=10, pady=10)
        report_frame.pack(fill='x', pady=5)
        
        tk.Button(report_frame, text="Generar Reporte Unificado Consolidado", 
                 command=self.generate_unified_consolidated_report, bg="lightblue", font=("Arial", 10, "bold")).pack(pady=2, fill='x')
        
        # Student Consultation Section
        student_frame = tk.LabelFrame(main_frame, text="Consulta de Estudiantes", padx=10, pady=10)
        student_frame.pack(fill='x', pady=5)
        
        tk.Button(student_frame, text="Consultar Estudiante", 
                 command=self.start_student_consultation, bg="lightyellow").pack(pady=2, fill='x')
        
        # Impact Analysis Section
        impact_frame = tk.LabelFrame(main_frame, text="Análisis de Impacto", padx=10, pady=10)
        impact_frame.pack(fill='x', pady=5)
        
        tk.Button(impact_frame, text="Generar Análisis de Beneficiarios", 
                 command=self.generate_beneficiary_analysis, bg="lightcoral").pack(pady=2, fill='x')
        
        # Document Generation Section
        doc_frame = tk.LabelFrame(main_frame, text="Generación de Oficios", padx=10, pady=10)
        doc_frame.pack(fill='x', pady=5)
        
        tk.Button(doc_frame, text="Generar Oficios Institucionales", 
                 command=self.generate_official_documents, bg="lightpink").pack(pady=2, fill='x')
        
        # Status Label
        self.status_label = tk.Label(self.root, text="", fg="blue", wraplength=500)
        self.status_label.pack(pady=10)

    def select_folder(self):
        nuevo_directorio = filedialog.askdirectory(initialdir=self.directorio_destino)
        if nuevo_directorio:
            self.directorio_destino = nuevo_directorio
            self.folder_label.config(text=f"Carpeta: {self.directorio_destino}")

    def update_status(self, message, color="blue"):
        self.status_label.config(text=message, fg=color)
        self.root.update()

    def validate_period(self, periodo):
        try:
            year, quarter = periodo.split('-')
            return len(year) == 4 and quarter.isdigit() and 1 <= int(quarter) <= 4
        except:
            return False

    def convert_timestamp_to_period(self, timestamp_value):
        """Convierte timestamp a formato YYYY-Q"""
        try:
            if pd.isna(timestamp_value):
                return None
            
            # Si ya es string en formato correcto, devolverlo
            if isinstance(timestamp_value, str) and '-' in timestamp_value:
                return timestamp_value
            
            # Convertir timestamp a datetime
            if isinstance(timestamp_value, pd.Timestamp):
                dt = timestamp_value
            else:
                dt = pd.to_datetime(timestamp_value)
            
            # Determinar el trimestre basado en el mes
            year = dt.year
            month = dt.month
            
            if month <= 3:
                quarter = 1
            elif month <= 6:
                quarter = 2
            elif month <= 9:
                quarter = 3
            else:
                quarter = 4
                
            return f"{year}-{quarter}"
        except:
            return None

    def verify_password(self):
        """Verifica la contraseña desde la hoja de cálculo pública"""
        try:
            df_password = pd.read_csv(self.url_password)
            password_correcto = str(df_password.iloc[0]['Contraseña'])
            
            for _ in range(3):
                password = simpledialog.askstring("Inicio de sesión", "Por favor ingrese la contraseña:", show='*')
                if password is None:
                    return False
                if password == password_correcto:
                    return True
                else:
                    messagebox.showerror("Error", "Contraseña incorrecta.")
            
            messagebox.showwarning("Acceso denegado", "Se han agotado los intentos.")
            return False
        except Exception as e:
            messagebox.showerror("Error", f"Error al verificar la contraseña: {e}")
            return False

    def load_data(self):
        """Carga todos los datos necesarios"""
        try:
            self.update_status("Cargando datos...")
            df_datos = pd.read_excel(self.url_datos)
            df_beneficiarios = pd.read_csv(self.url_beneficiarios)
            df_ubicacion = pd.read_excel(self.url_ubicacion)
            
            # NO normalizar nombres de columnas para mantener los nombres originales
            return df_datos, df_beneficiarios, df_ubicacion
        except Exception as e:
            self.update_status("Error cargando datos", "red")
            messagebox.showerror("Error", f"Error al cargar datos: {str(e)}")
            return None, None, None

    def generate_unified_consolidated_report(self):
        """Genera el reporte unificado consolidado con múltiples hojas - CORREGIDO CON FILTRADO POR PERÍODO"""
        periodo = self.period_entry.get()
        if not self.validate_period(periodo):
            messagebox.showerror("Error", "Formato de período inválido. Use el formato YYYY-Q")
            return
        
        try:
            df_datos, df_beneficiarios, df_ubicacion = self.load_data()
            if df_beneficiarios is None:
                return
            
            if not os.path.exists(self.directorio_destino):
                os.makedirs(self.directorio_destino)
            
            # Verificar columnas disponibles
            print(f"Columnas disponibles en beneficiarios: {list(df_beneficiarios.columns)}")
            print(f"Columnas disponibles en instituciones: {list(df_datos.columns)}")
            print(f"Buscando período: {periodo}")
            
            # Verificar que las columnas existen en beneficiarios
            if 'Período de registro' not in df_beneficiarios.columns:
                messagebox.showerror("Error", f"No se encontró la columna 'Período de registro' en los datos de beneficiarios.\nColumnas disponibles: {list(df_beneficiarios.columns)}")
                return
            
            if 'Centro de Educación' not in df_beneficiarios.columns:
                messagebox.showerror("Error", f"No se encontró la columna 'Centro de Educación' en los datos de beneficiarios.\nColumnas disponibles: {list(df_beneficiarios.columns)}")
                return
            
            if 'Qué voy a reportar' not in df_beneficiarios.columns:
                messagebox.showerror("Error", f"No se encontró la columna 'Qué voy a reportar' en los datos de beneficiarios.\nColumnas disponibles: {list(df_beneficiarios.columns)}")
                return
            
            # CORRECCIÓN PRINCIPAL: Filtrar instituciones por período
            if 'Período' in df_datos.columns:
                # Convertir la columna Período a formato YYYY-Q
                df_datos['Período_Convertido'] = df_datos['Período'].apply(self.convert_timestamp_to_period)
                print(f"Períodos únicos en instituciones: {df_datos['Período_Convertido'].unique()}")
                
                # Filtrar instituciones por período
                df_datos_filtrado = df_datos[df_datos['Período_Convertido'] == periodo]
                print(f"Instituciones encontradas para {periodo}: {len(df_datos_filtrado)}")
                
                if df_datos_filtrado.empty:
                    messagebox.showwarning("Sin datos de instituciones", f"No se encontraron instituciones para el período {periodo}")
                    # Continuar con df_datos original para no bloquear el proceso
                    df_datos_filtrado = df_datos
            else:
                print("⚠️ No se encontró columna 'Período' en instituciones, usando todas las instituciones")
                df_datos_filtrado = df_datos
            
            # Filtrar beneficiarios por período
            df_beneficiarios_filtrado = df_beneficiarios[df_beneficiarios['Período de registro'] == periodo]
            print(f"Registros de beneficiarios encontrados para {periodo}: {len(df_beneficiarios_filtrado)}")
            
            if df_beneficiarios_filtrado.empty:
                messagebox.showwarning("Sin datos", f"No se encontraron datos de beneficiarios para el período {periodo}")
                return
            
            # === GENERAR HOJA DE BENEFICIARIOS ===
            centros = pd.DataFrame({'Centro de Educación': df_beneficiarios_filtrado['Centro de Educación'].unique()})
            
            # Categories and their corresponding filters
            print("Valores únicos en 'Qué voy a reportar':")
            valores_reportar = df_beneficiarios_filtrado['Qué voy a reportar'].unique()
            for i, valor in enumerate(valores_reportar):
                print(f"  {i+1}. \"{valor}\"")
            
            categorias = {
                'Atenciones Individuales': 'Estudiante atendido Individualmente',
                'Asesorías': 'Asesorías a funcionarios',
                'Evaluaciones Psicopedagógicas': 'Beneficiarios de evaluación psicopedagógica',
                'DIAC': 'Beneficiarios DIAC o plan de intervención',
                'Capacitaciones Funcionarios': 'Capacitación a funcionario(s) (Docentes u otros)',
                'Sensibilizaciones': 'Sensibilización',
                'Capacitaciones Padres': 'Capacitación a padres de familia'
            }
            
            # Calculate counts for each category
            resultados = {}
            for nombre, filtro in categorias.items():
                data = df_beneficiarios_filtrado[df_beneficiarios_filtrado['Qué voy a reportar'] == filtro]
                print(f"Categoría '{nombre}' con filtro '{filtro}': {len(data)} registros")
                resultados[nombre] = centros['Centro de Educación'].map(
                    data['Centro de Educación'].value_counts()
                ).fillna(0).astype(int)
            
            # Calculate additional metrics
            sensibilizaciones = df_beneficiarios_filtrado[df_beneficiarios_filtrado['Qué voy a reportar'] == 'Sensibilización']
            
            # Verificar si existe la columna de participantes en sensibilización
            col_participantes_sensibilizacion = None
            for col in df_beneficiarios.columns:
                if 'participantes' in col.lower() and 'sensibilización' in col.lower():
                    col_participantes_sensibilizacion = col
                    break
            
            if col_participantes_sensibilizacion:
                print(f"Usando columna para participantes en sensibilización: '{col_participantes_sensibilizacion}'")
                personas_sensibilizadas = centros['Centro de Educación'].map(
                    sensibilizaciones.groupby('Centro de Educación')[col_participantes_sensibilizacion].sum()
                ).fillna(0).astype(int)
            else:
                print("⚠️ No se encontró columna de participantes en sensibilización, usando 0")
                personas_sensibilizadas = centros['Centro de Educación'].map(lambda x: 0)
            
            capacitaciones_padres = df_beneficiarios_filtrado[df_beneficiarios_filtrado['Qué voy a reportar'] == 'Capacitación a padres de familia']
            
            # Verificar si existe la columna de padres capacitados
            col_padres_capacitados = None
            for col in df_beneficiarios.columns:
                if 'padres' in col.lower() and 'capacitación' in col.lower() and 'número' in col.lower():
                    col_padres_capacitados = col
                    break
            
            if col_padres_capacitados:
                print(f"Usando columna para padres capacitados: '{col_padres_capacitados}'")
                padres_capacitados = centros['Centro de Educación'].map(
                    capacitaciones_padres.groupby('Centro de Educación')[col_padres_capacitados].sum()
                ).fillna(0).astype(int)
            else:
                print("⚠️ No se encontró columna de padres capacitados, usando 0")
                padres_capacitados = centros['Centro de Educación'].map(lambda x: 0)
            
            # Create Beneficiarios sheet
            beneficiarios = centros.copy()
            beneficiarios['Atenciones Individuales'] = resultados['Atenciones Individuales']
            beneficiarios['Asesorías'] = resultados['Asesorías']
            beneficiarios['Personas Sensibilizadas'] = personas_sensibilizadas
            beneficiarios['Padres y Cuidadores Capacitados'] = padres_capacitados
            beneficiarios['Evaluaciones Psicopedagógicas'] = resultados['Evaluaciones Psicopedagógicas']
            beneficiarios['DIAC'] = resultados['DIAC']
            
            # Create Actividades sheet
            actividades = centros.copy()
            actividades['Capacitaciones Funcionarios'] = resultados['Capacitaciones Funcionarios']
            actividades['Sensibilizaciones'] = resultados['Sensibilizaciones']
            actividades['Capacitaciones Padres'] = resultados['Capacitaciones Padres']
            
            # === GENERAR HOJA DE CENTROS FILTRADA POR PERÍODO ===
            print(f"Columnas en df_datos_filtrado: {list(df_datos_filtrado.columns)}")
            print(f"Columnas en df_ubicacion: {list(df_ubicacion.columns)}")
            
            # CORRECCIÓN: Usar df_datos_filtrado en lugar de df_datos
            df_centros = df_datos_filtrado.merge(df_ubicacion, how='left', 
                                               left_on='Nombre Corto de la Institución', 
                                               right_on='INSTITUCIÓN')
            
            print(f"✅ Centros después del filtrado por período: {len(df_centros)}")
            
            # Agregar enlaces de Google Maps si hay coordenadas
            if 'LATITUD' in df_centros.columns and 'LONGITUD' in df_centros.columns:
                df_centros['dirección_en_google_maps'] = "https://www.google.com/maps?q=" + \
                                                       df_centros['LATITUD'].astype(str) + "," + \
                                                       df_centros['LONGITUD'].astype(str)
            
            # === GUARDAR ARCHIVO CONSOLIDADO ===
            archivo_consolidado = os.path.join(self.directorio_destino, f"Reporte_Unificado_Consolidado_{periodo}.xlsx")
            
            with pd.ExcelWriter(archivo_consolidado, engine='openpyxl') as writer:
                # Hoja 1: Beneficiarios
                beneficiarios.to_excel(writer, sheet_name='Beneficiarios', index=False)
                
                # Hoja 2: Actividades  
                actividades.to_excel(writer, sheet_name='Actividades', index=False)
                
                # Hoja 3: Centros de Educación (AHORA FILTRADA POR PERÍODO)
                df_centros.to_excel(writer, sheet_name='Centros_Educacion', index=False)
            
            self.update_status(f"Reporte consolidado generado: {archivo_consolidado}", "green")
            messagebox.showinfo("Éxito", 
                              f"Reporte Unificado Consolidado generado exitosamente.\n\n"
                              f"Archivo: {archivo_consolidado}\n\n"
                              f"Registros procesados:\n"
                              f"• Beneficiarios: {len(df_beneficiarios_filtrado)}\n"
                              f"• Instituciones: {len(df_centros)}\n\n"
                              f"Hojas incluidas:\n"
                              f"• Beneficiarios (resumen por centro)\n"
                              f"• Actividades (actividades por centro)\n"
                              f"• Centros_Educacion (FILTRADA por período {periodo})")
            
        except Exception as e:
            self.update_status("Error generando reporte consolidado", "red")
            messagebox.showerror("Error", f"Error: {str(e)}")
            import traceback
            traceback.print_exc()

    def start_student_consultation(self):
        """Inicia el proceso de consulta de estudiantes"""
        if not self.verify_password():
            return
        
        self.student_consultation_window()

    def student_consultation_window(self):
        """Ventana para consulta de estudiantes"""
        try:
            # Cargar datos de estudiantes
            df_estudiantes = pd.read_excel(self.url_datos)
            df_ubicaciones = pd.read_excel(self.url_ubicacion)
            
            # Convertir la columna Período a formato YYYY-Q
            if 'Período' in df_estudiantes.columns:
                df_estudiantes['Período_Convertido'] = df_estudiantes['Período'].apply(self.convert_timestamp_to_period)
            
            # Usar los nombres exactos de las columnas
            if 'INSTITUCIÓN' in df_ubicaciones.columns and 'LATITUD' in df_ubicaciones.columns and 'LONGITUD' in df_ubicaciones.columns:
                # Renombrar columnas para estandarizar
                df_ubicaciones = df_ubicaciones.rename(columns={
                    'INSTITUCIÓN': 'Nombre Corto de la Institución',
                    'LATITUD': 'Latitud',
                    'LONGITUD': 'Longitud'
                })
                
                # Convertir a string las columnas de institución
                if 'Nombre Corto de la Institución' in df_estudiantes.columns:
                    df_estudiantes['Nombre Corto de la Institución'] = df_estudiantes['Nombre Corto de la Institución'].astype(str)
                if 'Nombre Corto de la Institución' in df_ubicaciones.columns:
                    df_ubicaciones['Nombre Corto de la Institución'] = df_ubicaciones['Nombre Corto de la Institución'].astype(str)
                
                # Merge de los dataframes
                df_combinado = pd.merge(
                    df_estudiantes,
                    df_ubicaciones[['Nombre Corto de la Institución', 'Latitud', 'Longitud']],
                    on='Nombre Corto de la Institución',
                    how='left'
                )
            else:
                # Continuar sin datos de ubicación
                df_combinado = df_estudiantes.copy()
                df_combinado['Latitud'] = None
                df_combinado['Longitud'] = None
            
            self.show_student_search_dialog(df_combinado)
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print("Error en student_consultation_window:", error_details)
            messagebox.showerror("Error", f"Error al cargar datos de estudiantes: {e}")

    def show_student_search_dialog(self, df_estudiantes):
        """Muestra el diálogo de búsqueda de estudiantes"""
        search_window = tk.Toplevel(self.root)
        search_window.title("Consulta de Estudiantes")
        search_window.geometry("500x300")
        
        # Frame principal
        main_frame = tk.Frame(search_window)
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Título
        tk.Label(main_frame, text="Consulta de Estudiantes", font=("Arial", 14, "bold")).pack(pady=5)
        
        # Período
        period_frame = tk.Frame(main_frame)
        period_frame.pack(pady=5)
        tk.Label(period_frame, text="Período (opcional):").pack(side=tk.LEFT)
        period_entry = tk.Entry(period_frame, width=15)
        period_entry.pack(side=tk.LEFT, padx=5)
        tk.Label(period_frame, text="(formato: 2025-1)").pack(side=tk.LEFT)
        
        # Nombre del estudiante
        name_frame = tk.Frame(main_frame)
        name_frame.pack(pady=10)
        tk.Label(name_frame, text="Nombre del estudiante:", font=("Arial", 12)).pack()
        search_entry = tk.Entry(name_frame, width=40, font=("Arial", 10))
        search_entry.pack(pady=5)
        
        # Resultado
        result_frame = tk.Frame(main_frame)
        result_frame.pack(fill='both', expand=True, pady=10)
        
        def buscar_estudiante():
            nombre = search_entry.get().strip()
            periodo_filtro = period_entry.get().strip()
            
            if not nombre:
                messagebox.showwarning("Advertencia", "Por favor ingrese un nombre")
                return
            
            # Filtrar por período si se especifica
            df_filtrado = df_estudiantes.copy()
            if periodo_filtro:
                if 'Período_Convertido' in df_filtrado.columns:
                    df_filtrado = df_filtrado[df_filtrado['Período_Convertido'] == periodo_filtro]
                else:
                    messagebox.showwarning("Advertencia", "No se pudo filtrar por período")
            
            resultado = self.search_student(df_filtrado, nombre)
            if resultado is not None and not resultado.empty:
                search_window.destroy()
                self.show_student_info(resultado)
        
        # Botón de búsqueda
        tk.Button(main_frame, text="Buscar Estudiante", command=buscar_estudiante, 
                 bg="lightblue", font=("Arial", 10, "bold")).pack(pady=10)
        
        search_entry.focus()
        search_entry.bind('<Return>', lambda event: buscar_estudiante())

    def search_student(self, df_estudiantes, nombre_busqueda):
        """Busca un estudiante por nombre"""
        try:
            if 'Apellidos y Nombres del Estudiante' not in df_estudiantes.columns:
                messagebox.showerror("Error", "No se encontró la columna de nombres de estudiantes")
                return None
            
            # Convertir a string y limpiar
            df_estudiantes['Apellidos y Nombres del Estudiante'] = df_estudiantes['Apellidos y Nombres del Estudiante'].astype(str)
            
            # Búsqueda exacta primero (case insensitive)
            coincidencias_exactas = df_estudiantes[
                df_estudiantes['Apellidos y Nombres del Estudiante'].str.contains(
                    nombre_busqueda, case=False, na=False, regex=False
                )
            ]
            
            if not coincidencias_exactas.empty:
                coincidencias = coincidencias_exactas
            else:
                # Búsqueda difusa usando thefuzz
                nombres = df_estudiantes['Apellidos y Nombres del Estudiante'].dropna().tolist()
                matches = process.extract(nombre_busqueda, nombres, limit=5)
                
                # Filtrar matches con score > 60
                good_matches = [match for match in matches if match[1] > 60]
                
                if good_matches:
                    nombres_coincidentes = [match[0] for match in good_matches]
                    coincidencias = df_estudiantes[
                        df_estudiantes['Apellidos y Nombres del Estudiante'].isin(nombres_coincidentes)
                    ]
                else:
                    coincidencias = pd.DataFrame()
            
            if coincidencias.empty:
                messagebox.showinfo("Sin coincidencias", "No se encontraron estudiantes con ese nombre.")
                return None
            
            if len(coincidencias) > 1:
                # Mostrar opciones para seleccionar
                opciones = []
                for idx, row in coincidencias.iterrows():
                    nombre = row['Apellidos y Nombres del Estudiante']
                    periodo = row.get('Período_Convertido', 'N/A')
                    institucion = row.get('Nombre Completo de la Institución', 'N/A')
                    opciones.append(f"{nombre} - {periodo} - {institucion}")
                
                mensaje = "Se encontraron varios estudiantes:\n" + "\n".join([f"{i+1}. {op}" for i, op in enumerate(opciones)])
                seleccion = simpledialog.askinteger("Selección", mensaje + "\n\nSeleccione el número:")
                if seleccion and 1 <= seleccion <= len(opciones):
                    return coincidencias.iloc[[seleccion - 1]]
                return None
            else:
                return coincidencias
                    
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print("Error en search_student:", error_details)
            messagebox.showerror("Error", f"Error en la búsqueda: {e}")
            return None


    def show_student_info(self, estudiante):
        """Muestra la información del estudiante"""
        if estudiante is None or estudiante.empty:
            return
        
        try:
            columnas_mostrar = [
                'Período', 'Período_Convertido', 'Apellidos y Nombres del Estudiante', 'Cédula',
                'Teléfono_Estudiante', 'Horas de práctica', 'Proyecto', 'Supervisor',
                'Servicio Social Educativo', 'Nombre Completo de la Institución',
                'Rector(a) o Autoridad ', 'Número total de estudiantes o beneficiarios de la institución',
                'Número total de docentes o funcionarios de la institución', 'Latitud', 'Longitud'
            ]
            
            # Construir información solo con columnas que existan
            info_lines = []
            row = estudiante.iloc[0]
            
            for col in columnas_mostrar:
                if col in estudiante.columns:
                    valor = row[col]
                    # Convertir a string y manejar valores nulos/NaN
                    if pd.isna(valor) or valor is None:
                        valor_str = 'No disponible'
                    else:
                        # Manejar diferentes tipos de datos
                        if isinstance(valor, (pd.Timestamp, pd.DatetimeIndex)):
                            valor_str = valor.strftime('%Y-%m-%d %H:%M:%S') if hasattr(valor, 'strftime') else str(valor)
                        else:
                            valor_str = str(valor)
                    
                    info_lines.append(f"{col}: {valor_str}")
            
            info = "\n".join(info_lines)
            
            # Agregar enlace a Google Maps si hay coordenadas disponibles
            try:
                if ('Latitud' in estudiante.columns and 'Longitud' in estudiante.columns and 
                    pd.notna(row['Latitud']) and pd.notna(row['Longitud'])):
                    lat = float(row['Latitud'])
                    lon = float(row['Longitud'])
                    maps_link = f"\n\nUbicación en Google Maps:\nhttps://www.google.com/maps?q={lat},{lon}"
                    info += maps_link
            except (ValueError, TypeError):
                # Si no se pueden convertir las coordenadas, simplemente no agregar el enlace
                pass
            
            # Crear ventana con scroll para la información
            info_window = tk.Toplevel(self.root)
            info_window.title("Información del Estudiante")
            info_window.geometry("700x500")
            
            # Frame para el contenido
            main_frame = tk.Frame(info_window)
            main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # Text widget con scrollbar
            text_widget = tk.Text(main_frame, wrap=tk.WORD, font=("Consolas", 10))
            scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=text_widget.yview)
            text_widget.configure(yscrollcommand=scrollbar.set)
            
            text_widget.insert(tk.END, info)
            text_widget.config(state=tk.DISABLED)
            
            text_widget.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Botón para cerrar
            tk.Button(info_window, text="Cerrar", command=info_window.destroy).pack(pady=5)
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print("Error mostrando información:", error_details)
            messagebox.showerror("Error", f"Error mostrando información del estudiante: {e}")

    def generate_beneficiary_analysis(self):
        """Genera análisis de beneficiarios"""
        periodo = self.period_entry.get()
        if not periodo:
            messagebox.showerror("Error", "Por favor ingrese un período")
            return
        
        try:
            self.update_status("Generando análisis de beneficiarios...")
            
            # Leer datos de encuesta
            data = pd.read_csv(self.url_encuesta)
            data.columns = data.columns.str.strip()
            
            # Filtrar por período
            data = data[data['Periodo'] == periodo]
            
            if data.empty:
                messagebox.showwarning("Sin datos", f"No se encontraron datos para el período {periodo}")
                return
            
            # Crear carpeta de salida
            output_dir = os.path.join(self.directorio_destino, "Beneficiarios")
            os.makedirs(output_dir, exist_ok=True)
            
            # Guardar datos en Excel
            excel_filename = os.path.join(output_dir, f'Encuesta_Beneficiarios_{periodo}.xlsx')
            data.to_excel(excel_filename, index=False)
            
            # Generar gráficos
            self.generate_beneficiary_charts(data, output_dir, periodo)
            
            # Generar documento Word
            self.create_beneficiary_document(output_dir, periodo)
            
            self.update_status(f"Análisis de beneficiarios completado en: {output_dir}", "green")
            messagebox.showinfo("Éxito", f"Análisis completado. Archivos guardados en:\n{output_dir}")
            
        except Exception as e:
            self.update_status("Error en análisis de beneficiarios", "red")
            messagebox.showerror("Error", f"Error: {str(e)}")

    def generate_beneficiary_charts(self, data, output_dir, periodo):
        """Genera gráficos para el análisis de beneficiarios"""
        plt.style.use('default')
        
        # 1. Gráfico de Ciudad
        if 'Ciudad - Institución' in data.columns:
            plt.figure(figsize=(10, 6))
            city_counts = data['Ciudad - Institución'].value_counts()
            city_counts.plot(kind='bar', color='skyblue')
            plt.title('Distribución de Respuestas por Ciudad')
            plt.xlabel('Ciudad')
            plt.ylabel('Número de Respuestas')
            plt.xticks(rotation=45, ha='right')
            for index, value in enumerate(city_counts):
                plt.text(index, value + 0.5, f'{value}\n({value / city_counts.sum() * 100:.1f}%)', 
                        ha='center', va='bottom')
            plt.tight_layout()
            plt.savefig(os.path.join(output_dir, f'ciudad_distribution_{periodo}.jpg'), dpi=300, bbox_inches='tight')
            plt.close()
        
        # 2. Gráfico de Institución
        if 'Institución' in data.columns:
            plt.figure(figsize=(12, 8))
            institution_counts = data['Institución'].value_counts()
            institution_counts.plot(kind='barh', color='lightgreen')
            plt.title('Distribución de Respuestas por Institución')
            plt.xlabel('Número de Respuestas')
            plt.ylabel('Institución')
            for index, value in enumerate(institution_counts):
                plt.text(value + 0.5, index, f'{value}\n({value / institution_counts.sum() * 100:.1f}%)', 
                        va='center')
            plt.tight_layout()
            plt.savefig(os.path.join(output_dir, f'institucion_distribution_{periodo}.jpg'), dpi=300, bbox_inches='tight')
            plt.close()
        
        # Generar nubes de palabras si existen las columnas correspondientes
        stop_words_spanish = {
            'el', 'la', 'los', 'las', 'un', 'una', 'de', 'del', 'a', 'al', 'que', 'en', 'por', 'para',
            'con', 'sin', 'sobre', 'y', 'o', 'pero', 'también', 'ni', 'como', 'muy', 'no', 'ha', 'su',
            'estudiante', 'estudiantes', 'niño', 'niña', 'niños', 'niñas', 'actividad', 'actividades'
        }
        
        # Nube de palabras para aspectos positivos
        pos_col = 'Comparta brevemente con nosotros aspectos positivos que ha experimentado u observado durante el proceso  de prácticas o vinculación'
        if pos_col in data.columns:
            positive_text = ' '.join(data[pos_col].dropna())
            if positive_text.strip():
                wordcloud = WordCloud(width=800, height=400, background_color='white', 
                                    stopwords=stop_words_spanish).generate(positive_text)
                wordcloud.to_file(os.path.join(output_dir, f'wordcloud_positive_{periodo}.png'))
        
        # Nube de palabras para aspectos a mejorar
        imp_col = 'Comparta con nosotros aspectos a mejorar en función de lo que ha experimentado u observado durante el proceso de prácticas o vinculación'
        if imp_col in data.columns:
            improvement_text = ' '.join(data[imp_col].dropna())
            if improvement_text.strip():
                wordcloud = WordCloud(width=800, height=400, background_color='white', 
                                    stopwords=stop_words_spanish).generate(improvement_text)
                wordcloud.to_file(os.path.join(output_dir, f'wordcloud_improvement_{periodo}.png'))

    def create_beneficiary_document(self, output_dir, periodo):
        """Crea documento Word con los gráficos"""
        doc = Document()
        doc.add_heading('Análisis de Encuesta - Beneficiarios', 0)
        doc.add_paragraph(f'Período: {periodo}')
        
        # Lista de imágenes a incluir
        images = [
            ('Distribución por Ciudad', f'ciudad_distribution_{periodo}.jpg'),
            ('Distribución por Institución', f'institucion_distribution_{periodo}.jpg'),
            ('Aspectos Positivos', f'wordcloud_positive_{periodo}.png'),
            ('Aspectos a Mejorar', f'wordcloud_improvement_{periodo}.png')
        ]
        
        for title, filename in images:
            filepath = os.path.join(output_dir, filename)
            if os.path.exists(filepath):
                doc.add_heading(title, level=1)
                doc.add_picture(filepath, width=Inches(6.0))
        
        # Guardar documento
        word_filename = os.path.join(output_dir, f'Beneficiarios_analisis_{periodo}.docx')
        doc.save(word_filename)

    def generate_official_documents(self):
        """Genera oficios institucionales - CORREGIDO PARA USAR SOLO INSTITUCIONES DEL PERÍODO"""
        periodo = self.period_entry.get()
        if not self.validate_period(periodo):
            messagebox.showerror("Error", "Formato de período inválido. Use el formato YYYY-Q")
            return
        
        # Verificar si existe el archivo plantilla
        plantilla_path = "Formato Oficio - Editable.docx"
        if not os.path.exists(plantilla_path):
            messagebox.showerror("Error", f"No se encontró el archivo plantilla: {plantilla_path}")
            return
        
        # Verificar si existe el reporte unificado consolidado
        reporte_path = os.path.join(self.directorio_destino, f"Reporte_Unificado_Consolidado_{periodo}.xlsx")
        if not os.path.exists(reporte_path):
            respuesta = messagebox.askyesno("Reporte no encontrado", 
                                          f"No se encontró el reporte unificado consolidado para {periodo}.\n¿Desea generarlo primero?")
            if respuesta:
                self.generate_unified_consolidated_report()
                if not os.path.exists(reporte_path):
                    return
            else:
                return
        
        try:
            self.update_status("Generando oficios institucionales...")
            
            # Cargar datos
            df_datos, _, _ = self.load_data()
            if df_datos is None:
                return
            
            # CORRECCIÓN PRINCIPAL: Filtrar instituciones por período ANTES de generar oficios
            if 'Período' in df_datos.columns:
                # Convertir la columna Período a formato YYYY-Q
                df_datos['Período_Convertido'] = df_datos['Período'].apply(self.convert_timestamp_to_period)
                print(f"Períodos únicos en instituciones: {df_datos['Período_Convertido'].unique()}")
                
                # Filtrar instituciones por período
                df_datos_filtrado = df_datos[df_datos['Período_Convertido'] == periodo]
                print(f"✅ Instituciones filtradas para {periodo}: {len(df_datos_filtrado)}")
                
                if df_datos_filtrado.empty:
                    messagebox.showwarning("Sin instituciones", f"No se encontraron instituciones para el período {periodo}")
                    return
            else:
                print("⚠️ No se encontró columna 'Período' en instituciones, usando todas las instituciones")
                df_datos_filtrado = df_datos
            
            # Cargar reporte unificado consolidado
            beneficiarios_df = pd.read_excel(reporte_path, sheet_name="Beneficiarios")
            actividades_df = pd.read_excel(reporte_path, sheet_name="Actividades")
            
            # Crear carpeta de oficios
            carpeta_oficios = os.path.join(self.directorio_destino, 'Oficios_Instituciones')
            os.makedirs(carpeta_oficios, exist_ok=True)
            
            # Opciones para generar oficios
            opcion = messagebox.askyesno("Generar Oficios", 
                                       f"¿Desea generar oficios para TODAS las instituciones del período {periodo}?\n\n" +
                                       f"Instituciones encontradas: {len(df_datos_filtrado)}\n\n" +
                                       "Sí = Todas las instituciones del período\n" + 
                                       "No = Seleccionar institución específica")
            
            if opcion:
                # Generar para todas las instituciones DEL PERÍODO
                self.generate_all_official_documents(df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo)
            else:
                # Generar para una institución específica DEL PERÍODO
                self.generate_single_official_document(df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo)
                
        except Exception as e:
            self.update_status("Error generando oficios", "red")
            messagebox.showerror("Error", f"Error: {str(e)}")

    def generate_all_official_documents(self, df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo):
        """Genera oficios para todas las instituciones DEL PERÍODO FILTRADO"""
        # CORRECCIÓN: Usar df_datos_filtrado en lugar de df_datos
        instituciones = df_datos_filtrado['Nombre Completo de la Institución'].dropna().unique()
        count = 0
        
        print(f"🏢 Generando oficios para {len(instituciones)} instituciones del período {periodo}")
        
        for idx, nombre_institucion in enumerate(instituciones, start=1):
            try:
                # Buscar coincidencia en beneficiarios
                nombre_coincidencia, score = process.extractOne(
                    nombre_institucion, 
                    beneficiarios_df['Centro de Educación'].dropna().tolist()
                )
                
                print(f"  {idx}. {nombre_institucion} → {nombre_coincidencia} (score: {score})")
                
                if score > 60:  # Solo procesar si hay buena coincidencia
                    self.create_official_document(
                        nombre_institucion, nombre_coincidencia, idx, 
                        df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo
                    )
                    count += 1
                else:
                    print(f"    ⚠️ Score muy bajo ({score}), omitiendo institución")
                    
            except Exception as e:
                print(f"❌ Error procesando {nombre_institucion}: {e}")
                continue
        
        self.update_status(f"Se generaron {count} oficios en: {carpeta_oficios}", "green")
        messagebox.showinfo("Éxito", f"Se generaron {count} oficios institucionales para el período {periodo}")

    def generate_single_official_document(self, df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo):
        """Genera oficio para una institución específica DEL PERÍODO FILTRADO"""
        # CORRECCIÓN: Usar df_datos_filtrado en lugar de df_datos
        instituciones = df_datos_filtrado['Nombre Completo de la Institución'].dropna().unique()
        
        # Crear ventana de selección
        selection_window = tk.Toplevel(self.root)
        selection_window.title(f"Seleccionar Institución - Período {periodo}")
        selection_window.geometry("500x400")
        
        tk.Label(selection_window, text=f"Seleccione una institución del período {periodo}:", font=("Arial", 12)).pack(pady=10)
        tk.Label(selection_window, text=f"Instituciones disponibles: {len(instituciones)}", font=("Arial", 10), fg="blue").pack()
        
        listbox = tk.Listbox(selection_window, height=15)
        scrollbar = ttk.Scrollbar(selection_window, orient="vertical", command=listbox.yview)
        listbox.configure(yscrollcommand=scrollbar.set)
        
        for institucion in sorted(instituciones):
            listbox.insert(tk.END, institucion)
        
        listbox.pack(side="left", fill="both", expand=True, padx=10)
        scrollbar.pack(side="right", fill="y")
        
        def generar_seleccionado():
            selection = listbox.curselection()
            if selection:
                nombre_institucion = listbox.get(selection[0])
                nombre_coincidencia, _ = process.extractOne(
                    nombre_institucion, 
                    beneficiarios_df['Centro de Educación'].dropna().tolist()
                )
                
                self.create_official_document(
                    nombre_institucion, nombre_coincidencia, 1,
                    df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo
                )
                
                selection_window.destroy()
                self.update_status(f"Oficio generado para: {nombre_institucion}", "green")
                messagebox.showinfo("Éxito", f"Oficio generado para: {nombre_institucion}")
            else:
                messagebox.showwarning("Selección", "Por favor seleccione una institución")
        
        tk.Button(selection_window, text="Generar Oficio", command=generar_seleccionado).pack(pady=10)

    def replace_text_in_word_document(self, doc, replacements):
        """Función robusta para reemplazar texto en documentos Word"""
        print("🔄 Iniciando reemplazos en documento Word...")
        
        # Función auxiliar para reemplazar en runs de un párrafo
        def replace_in_paragraph(paragraph, old_text, new_text):
            """Reemplaza texto en un párrafo manejando runs fragmentados"""
            if old_text in paragraph.text:
                # Obtener el texto completo del párrafo
                full_text = paragraph.text
                
                # Si contiene el marcador, reemplazar
                if old_text in full_text:
                    new_full_text = full_text.replace(old_text, new_text)
                    
                    # Limpiar todos los runs existentes
                    for run in paragraph.runs:
                        run.text = ""
                    
                    # Crear un nuevo run con el texto reemplazado
                    if paragraph.runs:
                        paragraph.runs[0].text = new_full_text
                    else:
                        paragraph.add_run(new_full_text)
                    
                    print(f"  ✅ Reemplazado '{old_text}' → '{new_text}' en párrafo")
                    return True
            return False
        
        # Función auxiliar para reemplazar en celdas de tabla
        def replace_in_cell(cell, old_text, new_text):
            """Reemplaza texto en una celda de tabla"""
            replaced = False
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, old_text, new_text):
                    replaced = True
            return replaced
        
        # Contador de reemplazos realizados
        total_replacements = 0
        
        # Realizar reemplazos
        for old_text, new_text in replacements.items():
            replacements_for_this_marker = 0
            
            # 1. Reemplazar en párrafos principales
            for paragraph in doc.paragraphs:
                if replace_in_paragraph(paragraph, old_text, new_text):
                    replacements_for_this_marker += 1
            
            # 2. Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if replace_in_cell(cell, old_text, new_text):
                            replacements_for_this_marker += 1
            
            # 3. Reemplazar en headers y footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if replace_in_paragraph(paragraph, old_text, new_text):
                            replacements_for_this_marker += 1
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if replace_in_paragraph(paragraph, old_text, new_text):
                            replacements_for_this_marker += 1
            
            total_replacements += replacements_for_this_marker
            
            if replacements_for_this_marker > 0:
                print(f"  ✅ '{old_text}' → '{new_text}' ({replacements_for_this_marker} reemplazos)")
            else:
                print(f"  ⚠️ '{old_text}' no encontrado en el documento")
        
        print(f"🎯 Total de reemplazos realizados: {total_replacements}")
        return total_replacements

    def create_official_document(self, nombre_institucion, nombre_coincidencia, numero_oficio, 
                               df_datos_filtrado, beneficiarios_df, actividades_df, carpeta_oficios, periodo):
        """Crea un documento oficial individual - USANDO DATOS FILTRADOS POR PERÍODO"""
        try:
            print(f"\n🏢 Generando oficio para: {nombre_institucion}")
            print(f"📊 Coincidencia en beneficiarios: {nombre_coincidencia}")
            
            # CORRECCIÓN: Usar df_datos_filtrado en lugar de df_datos
            row_institucion = df_datos_filtrado[df_datos_filtrado['Nombre Completo de la Institución'] == nombre_institucion].iloc[0]
            row_beneficiarios = beneficiarios_df[beneficiarios_df['Centro de Educación'] == nombre_coincidencia].iloc[0]
            row_actividades = actividades_df[actividades_df['Centro de Educación'] == nombre_coincidencia].iloc[0]
            
            print(f"📋 Datos de beneficiarios encontrados: {dict(row_beneficiarios)}")
            print(f"📋 Datos de actividades encontrados: {dict(row_actividades)}")
            
            # Cargar plantilla
            plantilla_path = "Formato Oficio - Editable.docx"
            nuevo_doc = Document(plantilla_path)
            
            # Buscar la columna correcta del representante
            col_representante = None
            for col in df_datos_filtrado.columns:
                if 'rector' in col.lower() and 'autoridad' in col.lower():
                    col_representante = col
                    print(f"✅ Encontrada columna de representante: '{col}'")
                    break
            
            if not col_representante:
                print("❌ No se encontró columna de representante")
                nombre_representante = "No disponible"
            else:
                nombre_representante = str(row_institucion.get(col_representante, "")).strip()
                print(f"👤 Nombre del representante obtenido: '{nombre_representante}'")
            
            # Obtener supervisor desde la columna "Supervisor"
            supervisor_proyecto = "Mg. Guillermo Andrade"  # Valor por defecto
            if 'Supervisor' in df_datos_filtrado.columns:
                supervisor_desde_bd = str(row_institucion.get('Supervisor', "")).strip()
                if supervisor_desde_bd and supervisor_desde_bd.lower() not in ['nan', 'none', '']:
                    supervisor_proyecto = supervisor_desde_bd
                    print(f"👨‍🏫 Supervisor obtenido desde BD: '{supervisor_proyecto}'")
                else:
                    print(f"⚠️ Supervisor vacío en BD, usando valor por defecto: '{supervisor_proyecto}'")
            else:
                print(f"❌ No se encontró columna 'Supervisor', usando valor por defecto: '{supervisor_proyecto}'")
            
            # Definir reemplazos usando los nombres correctos de columnas
            reemplazos = {
                # Limpiar marcador duplicado en el número de oficio
                "[Número de Oficio]": "",
                
                # Datos de la institución desde Google Sheets
                "[Título del representante de la institución]": str(row_institucion.get("Título del Rector o Autoridad", "")).strip(),
                "[Nombre del Representante de la institucion]": nombre_representante,
                "[Cargo del representante]": str(row_institucion.get("Cargo", "")).strip(),
                "[Nombre Completo de la Institución]": str(row_institucion.get("Nombre Completo de la Institución", "")).title(),
                
                # Datos del reporte consolidado
                "[Número de Capacitaciones Funcionarios]": str(int(row_actividades.get("Capacitaciones Funcionarios", 0))),
                "[Número de Funcionarios Capacitados]": str(int(row_actividades.get("Capacitaciones Funcionarios", 0))),
                "[Número de Capacitaciones Padres]": str(int(row_actividades.get("Capacitaciones Padres", 0))),
                "[Número de Padres Capacitados]": str(int(row_beneficiarios.get("Padres y Cuidadores Capacitados", 0))),
                "[Número de Sensibilizaciones]": str(int(row_actividades.get("Sensibilizaciones", 0))),
                "[Número de Personas Sensibilizadas]": str(int(row_beneficiarios.get("Personas Sensibilizadas", 0))),
                "[Número de Asesorías]": str(int(row_beneficiarios.get("Asesorías", 0))),
                "[Atenciones Individuales]": str(int(row_beneficiarios.get("Atenciones Individuales", 0))),
                "[Estudiantes DIAC]": str(int(row_beneficiarios.get("DIAC", 0))),
                "[Evaluaciones Psicopedagógicas]": str(int(row_beneficiarios.get("Evaluaciones Psicopedagógicas", 0))),
                
                # Datos adicionales
                "[Fecha]": date.today().strftime("%d de %B de %Y"),
                
                # Datos del proyecto
                "[Proyecto]": str(row_institucion.get("Proyecto", "Espacios de Apoyo Pedagógico Inclusivo")),
                "[Supervisor del Proyecto]": supervisor_proyecto,
            }
            
            # Mostrar información de depuración
            print(f"🔄 Reemplazos definidos para {nombre_institucion}:")
            for marcador, valor in reemplazos.items():
                print(f"  {marcador} → '{valor}'")
            
            # Usar la función de reemplazo robusta
            total_reemplazos = self.replace_text_in_word_document(nuevo_doc, reemplazos)
            
            # Generar nombre de archivo
            nombre_archivo = self.clean_filename(nombre_institucion)
            output_filename = f"Oficio_{nombre_archivo.replace(' ', '_')}_No_{numero_oficio:05d}_{periodo}.docx"
            output_filepath = os.path.join(carpeta_oficios, output_filename)
            
            # Guardar documento
            nuevo_doc.save(output_filepath)
            print(f"✅ Documento generado exitosamente: {output_filepath}")
            print(f"📊 Total de reemplazos realizados: {total_reemplazos}")
            
        except Exception as e:
            print(f"❌ Error generando documento para {nombre_institucion}: {e}")
            import traceback
            traceback.print_exc()
            raise

    def clean_filename(self, nombre):
        """Limpia el nombre para usarlo como nombre de archivo"""
        return re.sub(r'[<>:"/\\|?*]', '', nombre)

def main():
    root = tk.Tk()
    app = UnifiedReportApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()

